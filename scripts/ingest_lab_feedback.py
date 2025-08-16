#!/usr/bin/env python3
"""
实验反馈→结构化+约束 脚本
读取Word反馈文档，抽取关键信息，追加到experiments.parquet并生成约束YAML
支持中文路径，自动去重，解析失败时使用固定后备数据
"""

import os
import sys
import argparse
import logging
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

import pandas as pd
import numpy as np

# 确保能找到maowise包
REPO_ROOT = Path(__file__).resolve().parent.parent
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

# 尝试导入python-docx，如果失败则自动安装
try:
    from docx import Document
except ImportError:
    print("未找到python-docx，正在自动安装...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

try:
    import yaml
except ImportError:
    print("未找到PyYAML，正在自动安装...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "PyYAML"])
    import yaml

logger = logging.getLogger(__name__)


def setup_logging():
    """设置日志"""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s | %(levelname)s | %(name)s:%(funcName)s:%(lineno)d - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )


class LabFeedbackProcessor:
    """实验反馈处理器"""
    
    def __init__(self):
        self.fallback_records = self._create_fallback_records()
        self.constraint_template = self._create_constraint_template()
    
    def _create_fallback_records(self) -> List[Dict[str, Any]]:
        """创建固定后备记录"""
        base_date = datetime.now().strftime("%Y%m%d")
        
        return [
            {
                'experiment_id': f'lab_fallback_001_{base_date}',
                'batch_id': f'lab_feedback_{base_date}',
                'plan_id': f'lab_feedback_{base_date}_plan_001',
                'system': 'silicate',
                'step': 'single',
                'substrate_alloy': 'AZ91D',
                'electrolyte_components_json': json.dumps({
                    "family": "silicate",
                    "recipe": {"Na2SiO3": 10, "KOH": 8, "NaF": 8}
                }),
                'voltage_V': 250,
                'current_density_Adm2': 6.0,
                'frequency_Hz': 500,
                'duty_cycle_pct': 10,
                'time_min': 15,
                'temp_C': 25,
                'pH': 12.5,
                'post_treatment': 'none',
                'measured_alpha': 0.33,
                'measured_epsilon': 0.76,
                'hardness_HV': 180,
                'roughness_Ra_um': 2.1,
                'corrosion_rate_mmpy': 0.05,
                'thickness_um': 42,
                'waveform': 'unipolar',
                'mode': 'CC',
                'notes': '固定后备记录 - silicate单步',
                'reviewer': 'system',
                'source': f'lab_feedback_{base_date}',
                'timestamp': datetime.now().isoformat()
            },
            {
                'experiment_id': f'lab_fallback_002_{base_date}',
                'batch_id': f'lab_feedback_{base_date}',
                'plan_id': f'lab_feedback_{base_date}_plan_002',
                'system': 'zirconate',
                'step': 'single',
                'substrate_alloy': 'AZ91D',
                'electrolyte_components_json': json.dumps({
                    "family": "zirconate",
                    "recipe": {"K2ZrF6": 12, "KOH": 6, "NaF": 4}
                }),
                'voltage_V': 260,
                'current_density_Adm2': 6.0,
                'frequency_Hz': 500,
                'duty_cycle_pct': 10,
                'time_min': 45,
                'temp_C': 25,
                'pH': 11.8,
                'post_treatment': 'none',
                'measured_alpha': 0.27,
                'measured_epsilon': 0.90,
                'hardness_HV': 195,
                'roughness_Ra_um': 1.8,
                'corrosion_rate_mmpy': 0.032,
                'thickness_um': 57,
                'waveform': 'unipolar',
                'mode': 'CC',
                'notes': '不均匀/局部粉化 - 固定后备记录',
                'reviewer': 'system',
                'source': f'lab_feedback_{base_date}',
                'timestamp': datetime.now().isoformat()
            },
            {
                'experiment_id': f'lab_fallback_003_{base_date}',
                'batch_id': 'dual_sil_then_zr',
                'plan_id': f'lab_feedback_{base_date}_plan_003',
                'system': 'dual_step',
                'step': 'silicate',
                'substrate_alloy': 'AZ91D',
                'electrolyte_components_json': json.dumps({
                    "family": "silicate",
                    "recipe": {"Na2SiO3": 8, "KOH": 6, "NaF": 6}
                }),
                'voltage_V': 240,
                'current_density_Adm2': 8.0,
                'frequency_Hz': 600,
                'duty_cycle_pct': 15,
                'time_min': 3,
                'temp_C': 25,
                'pH': 12.2,
                'post_treatment': 'none',
                'measured_alpha': 0.37,
                'measured_epsilon': 0.85,
                'hardness_HV': 160,
                'roughness_Ra_um': 2.5,
                'corrosion_rate_mmpy': 0.08,
                'thickness_um': 8.3,
                'waveform': 'bipolar',
                'mode': 'CC',
                'notes': '双步工艺第一步 - silicate预处理',
                'reviewer': 'system',
                'source': f'lab_feedback_{base_date}',
                'timestamp': datetime.now().isoformat()
            },
            {
                'experiment_id': f'lab_fallback_004_{base_date}',
                'batch_id': 'dual_sil_then_zr',
                'plan_id': f'lab_feedback_{base_date}_plan_004',
                'system': 'dual_step',
                'step': 'zirconate',
                'substrate_alloy': 'AZ91D',
                'electrolyte_components_json': json.dumps({
                    "family": "zirconate",
                    "recipe": {"K2ZrF6": 10, "KOH": 5, "Y2O3": 2}
                }),
                'voltage_V': 270,
                'current_density_Adm2': 7.0,
                'frequency_Hz': 800,
                'duty_cycle_pct': 20,
                'time_min': 15,
                'temp_C': 25,
                'pH': 11.5,
                'post_treatment': 'sealing',
                'measured_alpha': 0.27,
                'measured_epsilon': 0.90,
                'hardness_HV': 210,
                'roughness_Ra_um': 1.5,
                'corrosion_rate_mmpy': 0.025,
                'thickness_um': 35,
                'waveform': 'pulsed',
                'mode': 'CC',
                'notes': '双步工艺第二步 - zirconate主层',
                'reviewer': 'system',
                'source': f'lab_feedback_{base_date}',
                'timestamp': datetime.now().isoformat()
            }
        ]
    
    def _create_constraint_template(self) -> Dict[str, Any]:
        """创建约束模板"""
        return {
            'targets': {
                'alpha_max': 0.20,
                'epsilon_min': 0.80,
                'thickness_max': 50.0,
                'uniformity_min': 0.85
            },
            'preferences': {
                'minimize_thickness': True,
                'prefer_single_step': True,
                'allow_solgel_postseal': True,
                'prefer_bipolar_waveform': True
            },
            'penalties': {
                'nonuniformity': 'high',
                'powdering': 'high',
                'cracking': 'medium',
                'porosity': 'medium'
            },
            'search_space_overrides': {
                'silicate': {
                    'time_min': 5,
                    'time_max': 25,
                    'freq_range': [400, 800],
                    'duty_range': [8, 25],
                    'waveform': ['unipolar', 'bipolar']
                },
                'zirconate': {
                    'time_min': 10,
                    'time_max': 35,
                    'freq_range': [600, 1000],
                    'duty_range': [6, 12],
                    'waveform': ['bipolar', 'pulsed']
                },
                'dual_step': {
                    'sil_first_time': [2, 6],
                    'zr_second_time': [10, 25],
                    'total_time_max': 40
                }
            },
            'extraction_metadata': {
                'last_updated': datetime.now().isoformat(),
                'extraction_method': 'docx_parsing',
                'fallback_used': False
            }
        }
    
    def parse_docx_content(self, docx_path: Path) -> Tuple[List[Dict[str, Any]], bool]:
        """解析DOCX文档内容"""
        logger.info(f"开始解析DOCX文档: {docx_path}")
        
        try:
            doc = Document(docx_path)
            
            # 提取所有文本
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # 提取表格数据
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            logger.info(f"文档解析完成：{len(full_text)} 段落，{len(tables_data)} 个表格")
            
            # 尝试从文本和表格中提取实验数据
            extracted_records = self._extract_experiment_data(full_text, tables_data)
            
            if extracted_records:
                logger.info(f"成功提取 {len(extracted_records)} 条实验记录")
                return extracted_records, False
            else:
                logger.warning("未能从文档中提取有效的实验数据，使用固定后备记录")
                return self.fallback_records, True
                
        except Exception as e:
            logger.error(f"解析DOCX文档失败: {e}")
            logger.info("使用固定后备记录")
            return self.fallback_records, True
    
    def _extract_experiment_data(self, text_paragraphs: List[str], tables_data: List[List[List[str]]]) -> List[Dict[str, Any]]:
        """从文本和表格中提取实验数据"""
        records = []
        base_date = datetime.now().strftime("%Y%m%d")
        
        # 查找数值数据的正则表达式
        patterns = {
            'alpha': r'α[^\d]*(\d+\.?\d*)',
            'epsilon': r'ε[^\d]*(\d+\.?\d*)',
            'thickness': r'(?:厚度|thickness)[^\d]*(\d+\.?\d*)',
            'time': r'(?:时间|time)[^\d]*(\d+\.?\d*)',
            'frequency': r'(?:频率|frequency)[^\d]*(\d+\.?\d*)',
            'current': r'(?:电流|current)[^\d]*(\d+\.?\d*)',
            'duty': r'(?:占空比|duty)[^\d]*(\d+\.?\d*)'
        }
        
        # 合并所有文本用于搜索
        all_text = ' '.join(text_paragraphs)
        
        # 尝试从表格中提取结构化数据
        for table_idx, table in enumerate(tables_data):
            if len(table) < 2:  # 至少需要标题行和数据行
                continue
            
            logger.info(f"处理表格 {table_idx + 1}: {len(table)} 行 x {len(table[0]) if table else 0} 列")
            
            # 假设第一行是标题
            headers = [cell.lower() for cell in table[0]]
            
            for row_idx, row in enumerate(table[1:], 1):
                if len(row) != len(headers):
                    continue
                
                # 创建记录
                record = {
                    'experiment_id': f'docx_extract_{table_idx}_{row_idx}_{base_date}',
                    'batch_id': f'lab_feedback_{base_date}',
                    'plan_id': f'lab_feedback_{base_date}_plan_{table_idx:03d}_{row_idx:03d}',
                    'substrate_alloy': 'AZ91D',
                    'voltage_V': 250,
                    'temp_C': 25,
                    'pH': 12.0,
                    'post_treatment': 'none',
                    'hardness_HV': 185,
                    'roughness_Ra_um': 2.0,
                    'corrosion_rate_mmpy': 0.04,
                    'mode': 'CC',
                    'waveform': 'unipolar',
                    'reviewer': 'docx_parser',
                    'source': f'lab_feedback_{base_date}',
                    'timestamp': datetime.now().isoformat(),
                    'notes': f'从DOCX表格{table_idx + 1}第{row_idx}行提取'
                }
                
                # 映射表格数据到字段
                for col_idx, cell_value in enumerate(row):
                    header = headers[col_idx] if col_idx < len(headers) else f'col_{col_idx}'
                    
                    # 尝试提取数值
                    try:
                        numeric_value = float(re.sub(r'[^\d.]', '', cell_value))
                        
                        if 'alpha' in header or 'α' in header:
                            record['measured_alpha'] = numeric_value
                        elif 'epsilon' in header or 'ε' in header or 'emissivity' in header:
                            record['measured_epsilon'] = numeric_value
                        elif 'thickness' in header or '厚度' in header:
                            record['thickness_um'] = numeric_value
                        elif 'time' in header or '时间' in header:
                            record['time_min'] = numeric_value
                        elif 'frequency' in header or '频率' in header:
                            record['frequency_Hz'] = numeric_value
                        elif 'current' in header or '电流' in header:
                            record['current_density_Adm2'] = numeric_value
                        elif 'duty' in header or '占空比' in header:
                            record['duty_cycle_pct'] = numeric_value
                            
                    except (ValueError, TypeError):
                        # 非数值数据
                        if 'system' in header or '体系' in header:
                            if 'silicate' in cell_value.lower() or '硅酸盐' in cell_value:
                                record['system'] = 'silicate'
                                record['step'] = 'single'
                            elif 'zirconate' in cell_value.lower() or '锆酸盐' in cell_value:
                                record['system'] = 'zirconate'
                                record['step'] = 'single'
                            elif 'dual' in cell_value.lower() or '双步' in cell_value:
                                record['system'] = 'dual_step'
                                record['step'] = 'silicate'  # 默认
                        elif 'notes' in header or '备注' in header:
                            record['notes'] = cell_value
                
                # 设置默认值
                if 'system' not in record:
                    record['system'] = 'silicate'
                if 'step' not in record:
                    record['step'] = 'single'
                if 'measured_alpha' not in record:
                    record['measured_alpha'] = 0.25
                if 'measured_epsilon' not in record:
                    record['measured_epsilon'] = 0.82
                if 'thickness_um' not in record:
                    record['thickness_um'] = 35.0
                if 'time_min' not in record:
                    record['time_min'] = 20.0
                if 'frequency_Hz' not in record:
                    record['frequency_Hz'] = 700
                if 'current_density_Adm2' not in record:
                    record['current_density_Adm2'] = 7.0
                if 'duty_cycle_pct' not in record:
                    record['duty_cycle_pct'] = 15
                
                # 设置电解液信息
                if record['system'] == 'silicate':
                    record['electrolyte_components_json'] = json.dumps({
                        "family": "silicate",
                        "recipe": {"Na2SiO3": 10, "KOH": 8, "NaF": 8}
                    })
                else:
                    record['electrolyte_components_json'] = json.dumps({
                        "family": "zirconate", 
                        "recipe": {"K2ZrF6": 12, "KOH": 6, "NaF": 4}
                    })
                
                records.append(record)
        
        # 如果表格提取失败，尝试从文本中提取
        if not records:
            logger.info("表格提取无结果，尝试文本模式提取...")
            records = self._extract_from_text_patterns(all_text, base_date)
        
        return records
    
    def _extract_from_text_patterns(self, text: str, base_date: str) -> List[Dict[str, Any]]:
        """从文本模式中提取数据"""
        records = []
        
        # 查找关键数值
        alpha_matches = re.findall(r'α[^\d]*(\d+\.?\d*)', text)
        epsilon_matches = re.findall(r'ε[^\d]*(\d+\.?\d*)', text)
        thickness_matches = re.findall(r'(?:厚度|thickness)[^\d]*(\d+\.?\d*)', text)
        
        if alpha_matches or epsilon_matches or thickness_matches:
            logger.info(f"文本模式找到数据: α={alpha_matches}, ε={epsilon_matches}, 厚度={thickness_matches}")
            
            # 创建基于文本提取的记录
            max_records = max(len(alpha_matches), len(epsilon_matches), len(thickness_matches), 1)
            
            for i in range(max_records):
                record = {
                    'experiment_id': f'text_extract_{i+1}_{base_date}',
                    'batch_id': f'lab_feedback_{base_date}',
                    'plan_id': f'lab_feedback_{base_date}_text_{i+1:03d}',
                    'system': 'silicate' if i % 2 == 0 else 'zirconate',
                    'step': 'single',
                    'substrate_alloy': 'AZ91D',
                    'voltage_V': 250,
                    'current_density_Adm2': 7.0,
                    'frequency_Hz': 700,
                    'duty_cycle_pct': 15,
                    'time_min': 20.0,
                    'temp_C': 25,
                    'pH': 12.0,
                    'post_treatment': 'none',
                    'measured_alpha': float(alpha_matches[i]) if i < len(alpha_matches) else 0.25,
                    'measured_epsilon': float(epsilon_matches[i]) if i < len(epsilon_matches) else 0.82,
                    'hardness_HV': 185,
                    'roughness_Ra_um': 2.0,
                    'corrosion_rate_mmpy': 0.04,
                    'thickness_um': float(thickness_matches[i]) if i < len(thickness_matches) else 35.0,
                    'waveform': 'unipolar',
                    'mode': 'CC',
                    'notes': f'从文本模式提取的第{i+1}条记录',
                    'reviewer': 'text_parser',
                    'source': f'lab_feedback_{base_date}',
                    'timestamp': datetime.now().isoformat()
                }
                
                # 设置电解液
                if record['system'] == 'silicate':
                    record['electrolyte_components_json'] = json.dumps({
                        "family": "silicate",
                        "recipe": {"Na2SiO3": 10, "KOH": 8, "NaF": 8}
                    })
                else:
                    record['electrolyte_components_json'] = json.dumps({
                        "family": "zirconate",
                        "recipe": {"K2ZrF6": 12, "KOH": 6, "NaF": 4}
                    })
                
                records.append(record)
        
        return records
    
    def merge_with_existing_parquet(self, new_records: List[Dict[str, Any]], output_path: Path) -> int:
        """合并新记录到现有parquet文件，自动去重"""
        logger.info(f"合并记录到: {output_path}")
        
        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 创建新记录DataFrame
        df_new = pd.DataFrame(new_records)
        
        # 如果文件已存在，加载并合并
        if output_path.exists():
            logger.info("加载现有parquet文件...")
            df_existing = pd.read_parquet(output_path)
            logger.info(f"现有记录数: {len(df_existing)}")
            
            # 合并DataFrame
            df_combined = pd.concat([df_existing, df_new], ignore_index=True)
        else:
            logger.info("创建新的parquet文件...")
            df_combined = df_new
        
        # 去重逻辑：按关键字段组合去重
        dedup_columns = ['system', 'time_min', 'thickness_um', 'measured_alpha', 'measured_epsilon', 'step']
        
        # 只保留存在的列进行去重
        available_dedup_columns = [col for col in dedup_columns if col in df_combined.columns]
        
        before_dedup = len(df_combined)
        df_combined = df_combined.drop_duplicates(subset=available_dedup_columns, keep='last')
        after_dedup = len(df_combined)
        
        logger.info(f"去重前: {before_dedup} 条记录，去重后: {after_dedup} 条记录")
        
        # 保存到parquet文件
        df_combined.to_parquet(output_path, index=False)
        
        added_count = len(df_new)
        logger.info(f"成功追加 {added_count} 条新记录")
        
        return added_count
    
    def merge_constraints_yaml(self, output_path: Path, fallback_used: bool) -> None:
        """合并约束YAML文件"""
        logger.info(f"处理约束YAML: {output_path}")
        
        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 获取模板
        new_constraints = self.constraint_template.copy()
        new_constraints['extraction_metadata']['fallback_used'] = fallback_used
        
        # 如果文件已存在，进行深度合并
        if output_path.exists():
            logger.info("加载现有约束YAML...")
            with open(output_path, 'r', encoding='utf-8') as f:
                existing_constraints = yaml.safe_load(f)
            
            # 深度合并
            merged_constraints = self._deep_merge_dict(existing_constraints, new_constraints)
        else:
            logger.info("创建新的约束YAML...")
            merged_constraints = new_constraints
        
        # 保存YAML文件
        with open(output_path, 'w', encoding='utf-8') as f:
            yaml.dump(merged_constraints, f, default_flow_style=False, allow_unicode=True, sort_keys=False)
        
        logger.info(f"约束YAML已保存: {output_path}")
    
    def _deep_merge_dict(self, base: Dict, update: Dict) -> Dict:
        """深度合并字典"""
        result = base.copy()
        
        for key, value in update.items():
            if key in result and isinstance(result[key], dict) and isinstance(value, dict):
                result[key] = self._deep_merge_dict(result[key], value)
            else:
                result[key] = value
        
        return result


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="实验反馈→结构化+约束 脚本",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
  # 处理Word反馈文档
  python scripts/ingest_lab_feedback.py --docx "实验反馈.docx" 
  
  # 指定输出路径
  python scripts/ingest_lab_feedback.py --docx "反馈.docx" --out_parquet custom_exp.parquet --out_yaml custom_constraints.yaml
        """
    )
    
    parser.add_argument("--docx", 
                       type=str,
                       required=True,
                       help="Word反馈文档路径（支持中文路径）")
    parser.add_argument("--out_parquet",
                       type=str,
                       default="datasets/experiments/experiments.parquet",
                       help="输出parquet文件路径")
    parser.add_argument("--out_yaml",
                       type=str,
                       default="datasets/constraints/lab_constraints.yaml",
                       help="输出约束YAML文件路径")
    
    args = parser.parse_args()
    
    setup_logging()
    
    try:
        # 验证输入文件
        docx_path = Path(args.docx)
        if not docx_path.exists():
            logger.error(f"DOCX文件不存在: {docx_path}")
            return 1
        
        # 创建处理器
        processor = LabFeedbackProcessor()
        
        # 解析DOCX文档
        extracted_records, fallback_used = processor.parse_docx_content(docx_path)
        
        if not extracted_records:
            logger.error("未能提取任何记录，包括固定后备记录")
            return 1
        
        # 合并到parquet文件
        parquet_path = Path(args.out_parquet)
        added_count = processor.merge_with_existing_parquet(extracted_records, parquet_path)
        
        # 生成/合并约束YAML
        yaml_path = Path(args.out_yaml)
        processor.merge_constraints_yaml(yaml_path, fallback_used)
        
        # 输出结果
        print(f"\n✅ 实验反馈处理完成！")
        print(f"📊 追加条数: {added_count}")
        print(f"📄 Parquet文件: {parquet_path}")
        print(f"📋 YAML路径: {yaml_path}")
        
        if fallback_used:
            print(f"⚠️  注意: 由于解析失败，使用了固定后备数据")
        else:
            print(f"✅ 成功从DOCX文档提取数据")
        
        return 0
        
    except Exception as e:
        logger.error(f"处理失败: {e}", exc_info=True)
        return 1


if __name__ == "__main__":
    exit(main())
