#!/usr/bin/env python3
"""
创建测试用的实验反馈Word文档
用于验证ingest_lab_feedback.py脚本的功能
"""

from docx import Document
from pathlib import Path

def create_test_docx():
    """创建测试用的DOCX文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('MAO工艺实验反馈报告', 0)
    
    # 添加概述段落
    doc.add_paragraph('本报告总结了最近进行的微弧氧化实验结果，包括硅酸盐和锆酸盐体系的性能数据。')
    
    # 添加表格数据
    doc.add_heading('实验数据汇总', level=1)
    
    # 创建数据表格
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    headers = ['体系', 'α值', 'ε值', '厚度(μm)', '时间(min)', '频率(Hz)', '电流(A/dm²)', '备注']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    # 添加数据行
    data_rows = [
        ['silicate', '0.28', '0.83', '38.5', '18', '650', '7.2', '均匀性良好'],
        ['zirconate', '0.22', '0.91', '52.3', '35', '800', '6.8', '轻微不均匀'],
        ['silicate', '0.31', '0.79', '35.2', '15', '550', '8.1', '局部粉化'],
        ['dual_step_sil', '0.35', '0.86', '12.1', '4', '700', '9.0', '双步第一阶段'],
        ['dual_step_zr', '0.24', '0.92', '41.8', '20', '850', '7.5', '双步第二阶段']
    ]
    
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value
    
    # 添加文本描述部分
    doc.add_heading('关键发现', level=1)
    findings = [
        '硅酸盐体系在α = 0.28-0.31, ε = 0.79-0.83范围内表现稳定',
        '锆酸盐体系实现了较低的α值（0.22）和较高的ε值（0.91）',
        '厚度控制在35-52μm之间，符合设计要求',
        '双步工艺显示出良好的分层效果',
        '部分样品存在不均匀和粉化问题，需要优化工艺参数'
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    # 添加工艺参数部分
    doc.add_heading('工艺参数详细信息', level=1)
    
    param_text = '''
    硅酸盐电解液配方：Na2SiO3 10 g/L, KOH 8 g/L, NaF 8 g/L
    锆酸盐电解液配方：K2ZrF6 12 g/L, KOH 6 g/L, NaF 4 g/L
    
    工艺条件：
    - 电压范围：240-270V
    - 温度：25±2°C  
    - pH值：11.5-12.5
    - 波形：单极脉冲，部分使用双极
    '''
    doc.add_paragraph(param_text)
    
    # 添加建议部分
    doc.add_heading('改进建议', level=1)
    suggestions = [
        '针对不均匀问题，建议降低电流密度至6.0 A/dm²',
        '粉化问题可通过调整频率至600-800Hz范围解决',
        '双步工艺第一步时间可优化至2-6分钟',
        '建议增加封孔处理以提高耐蚀性'
    ]
    
    for suggestion in suggestions:
        doc.add_paragraph(suggestion, style='List Number')
    
    return doc

def main():
    """主函数"""
    # 创建测试文档
    doc = create_test_docx()
    
    # 保存文档
    test_dir = Path("test_data")
    test_dir.mkdir(exist_ok=True)
    
    docx_path = test_dir / "实验反馈测试文档.docx"
    doc.save(docx_path)
    
    print(f"✅ 测试DOCX文档已创建: {docx_path}")
    print(f"📊 包含5条实验数据记录")
    print(f"📋 可用于测试 ingest_lab_feedback.py 脚本")

if __name__ == "__main__":
    main()
